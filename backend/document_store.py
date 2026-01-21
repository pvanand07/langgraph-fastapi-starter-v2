"""Document storage system with PDF parsing and search."""

import os
import re
import uuid
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional, Tuple
from sqlalchemy import create_engine, Column, Integer, String, Text, Index, and_, text
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session
import tiktoken
from config import DOCUMENTS_DB_PATH

try:
    import pymupdf4llm
    import pymupdf as fitz
except ImportError:
    raise RuntimeError("Please install pymupdf4llm: pip install pymupdf4llm")

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# Database setup
os.makedirs(os.path.dirname(DOCUMENTS_DB_PATH), exist_ok=True)
engine = create_engine(f"sqlite:///{DOCUMENTS_DB_PATH}", echo=False)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

# Token encoder
enc = tiktoken.get_encoding("cl100k_base")


class Page(Base):
    """Page table for storing document pages."""
    __tablename__ = 'page'
    
    id = Column(Integer, primary_key=True, autoincrement=True)
    user_id = Column(Text, nullable=False)
    doc_id = Column(Text, nullable=False)
    page_number = Column(Integer, nullable=False)
    content = Column(Text, nullable=False)
    token_count = Column(Integer, nullable=False)
    
    __table_args__ = (
        Index('ix_user_doc_page', 'user_id', 'doc_id', 'page_number', unique=True),
        Index('ix_user_id', 'user_id'),
        Index('ix_doc_id', 'doc_id'),
    )


class Document(Base):
    """Document metadata table."""
    __tablename__ = 'document'
    
    id = Column(Integer, primary_key=True, autoincrement=True)
    user_id = Column(Text, nullable=False)
    doc_id = Column(Text, nullable=False)
    doc_name = Column(Text, nullable=False)
    page_count = Column(Integer, nullable=False, default=0)
    created_at = Column(Text, nullable=False, default="")
    
    __table_args__ = (
        Index('ix_user_doc', 'user_id', 'doc_id', unique=True),
    )


def create_tables():
    """Create all database tables."""
    Base.metadata.create_all(bind=engine)
    
    # Create FTS5 virtual table for full-text search
    with engine.connect() as conn:
        result = conn.execute(text("""
            SELECT name FROM sqlite_master 
            WHERE type='table' AND name='page_fts'
        """)).fetchone()
        
        if not result:
            conn.execute(text("""
                CREATE VIRTUAL TABLE page_fts USING fts5(
                    content,
                    user_id UNINDEXED,
                    doc_id UNINDEXED,
                    page_number UNINDEXED
                )
            """))
            conn.commit()


def generate_doc_id(filename: str) -> str:
    """Generate document ID from filename."""
    name_without_ext = Path(filename).stem
    
    if len(name_without_ext) <= 20:
        doc_id = re.sub(r'[^a-zA-Z0-9]', '_', name_without_ext)
        doc_id = re.sub(r'_+', '_', doc_id).strip('_')
        return doc_id.lower()
    else:
        # Smart truncation for long names
        words = re.split(r'\s+', name_without_ext)
        processed_parts = []
        for word in words:
            if not word:
                continue
            if re.match(r'^[a-zA-Z0-9]+$', word):
                processed_parts.append(word[0].upper())
            else:
                parts = re.split(r'[^a-zA-Z0-9]+', word)
                for i, part in enumerate(parts):
                    if part:
                        if i > 0:
                            processed_parts.append('_')
                        processed_parts.append(part[0].upper())
        
        doc_id = ''.join(processed_parts)
        return doc_id.lower()[:20]  # Limit to 20 chars


def extract_pages_text_from_bytes(pdf_bytes: bytes, doc_id: str, user_id: str, filename: str = "") -> List[Dict]:
    """Extract text from PDF bytes and return list of page dictionaries."""
    pages = []
    
    try:
        # Use pymupdf4llm for better extraction
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Extract text with markdown formatting
            try:
                text = pymupdf4llm.to_markdown(page)
            except Exception:
                # Fallback to plain text
                text = page.get_text()
            
            if not text or not text.strip():
                continue
            
            # Calculate token count
            token_count = len(enc.encode(text))
            
            pages.append({
                "user_id": user_id,
                "doc_id": doc_id,
                "page_number": page_num + 1,
                "content": text,
                "token_count": token_count
            })
        
        doc.close()
    except Exception as e:
        raise RuntimeError(f"Error extracting PDF pages: {e}")
    
    return pages


def extract_pages_text_from_docx_bytes(docx_bytes: bytes, doc_id: str, user_id: str, filename: str = "") -> List[Dict]:
    """Extract text from DOCX bytes and return list of page dictionaries."""
    if DocxDocument is None:
        raise RuntimeError("Please install python-docx: pip install python-docx")
    
    pages = []
    
    try:
        from io import BytesIO
        doc = DocxDocument(BytesIO(docx_bytes))
        
        # DOCX doesn't have pages like PDF, so we'll split by paragraphs
        # Group paragraphs into "pages" of roughly similar size
        all_text_parts = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                all_text_parts.append(text)
        
        # Also extract text from tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_text.append(" | ".join(row_text))
            if table_text:
                all_text_parts.append("\n".join(table_text))
        
        if not all_text_parts:
            raise RuntimeError("No text content found in DOCX file")
        
        # Split into pages (approximately 2000 tokens per page)
        # Estimate tokens per character (roughly 4 chars per token)
        chars_per_page = 8000  # ~2000 tokens * 4 chars/token
        
        page_num = 1
        current_page_text = ""
        current_chars = 0
        
        for part in all_text_parts:
            part_chars = len(part)
            
            if current_chars + part_chars > chars_per_page and current_page_text:
                # Save current page
                token_count = len(enc.encode(current_page_text))
                pages.append({
                    "user_id": user_id,
                    "doc_id": doc_id,
                    "page_number": page_num,
                    "content": current_page_text,
                    "token_count": token_count
                })
                page_num += 1
                current_page_text = part
                current_chars = part_chars
            else:
                if current_page_text:
                    current_page_text += "\n\n" + part
                else:
                    current_page_text = part
                current_chars += part_chars + 2  # +2 for "\n\n"
        
        # Add the last page
        if current_page_text:
            token_count = len(enc.encode(current_page_text))
            pages.append({
                "user_id": user_id,
                "doc_id": doc_id,
                "page_number": page_num,
                "content": current_page_text,
                "token_count": token_count
            })
        
    except Exception as e:
        raise RuntimeError(f"Error extracting DOCX pages: {e}")
    
    return pages


def store_pages(pages: List[Dict], user_id: str, doc_id: str, doc_name: str):
    """Store pages and document metadata in SQLite."""
    create_tables()
    
    with SessionLocal() as db:
        try:
            # Delete existing pages and document for this user/doc_id
            db.query(Page).filter(
                and_(Page.user_id == user_id, Page.doc_id == doc_id)
            ).delete()
            
            db.query(Document).filter(
                and_(Document.user_id == user_id, Document.doc_id == doc_id)
            ).delete()
            
            # Delete from FTS table
            db.execute(text("""
                DELETE FROM page_fts 
                WHERE user_id = :user_id AND doc_id = :doc_id
            """), {"user_id": user_id, "doc_id": doc_id})
            
            db.commit()
            
            if pages:
                # Insert pages
                page_objects = [Page(**page) for page in pages]
                db.add_all(page_objects)
                db.flush()  # Get IDs
                
                # Insert into FTS table
                for page_obj in page_objects:
                    db.execute(text("""
                        INSERT INTO page_fts (content, user_id, doc_id, page_number)
                        VALUES (:content, :user_id, :doc_id, :page_number)
                    """), {
                        "content": page_obj.content,
                        "user_id": page_obj.user_id,
                        "doc_id": page_obj.doc_id,
                        "page_number": page_obj.page_number
                    })
                
                # Insert document metadata
                doc = Document(
                    user_id=user_id,
                    doc_id=doc_id,
                    doc_name=doc_name,
                    page_count=len(pages),
                    created_at=datetime.utcnow().isoformat()
                )
                db.add(doc)
                db.commit()
        except Exception as e:
            db.rollback()
            raise e


def list_documents(user_id: str) -> List[Dict]:
    """List all documents for a user."""
    create_tables()
    
    with SessionLocal() as db:
        docs = db.query(Document).filter(Document.user_id == user_id).all()
        return [
            {
                "doc_id": doc.doc_id,
                "doc_name": doc.doc_name,
                "user_id": doc.user_id,
                "page_count": doc.page_count,
                "created_at": doc.created_at
            }
            for doc in docs
        ]


def search_documents(user_id: str, query: str, doc_ids: Optional[List[str]] = None, limit: int = 10) -> List[Dict]:
    """
    Search documents using FTS5.
    
    Returns list of matching pages with doc_id, page_number, and content snippet.
    """
    create_tables()
    
    with SessionLocal() as db:
        # Build FTS query
        fts_query = f"""
            SELECT doc_id, page_number, content, 
                   snippet(page_fts, 0, '<mark>', '</mark>', '...', 32) as snippet
            FROM page_fts
            WHERE user_id = :user_id AND page_fts MATCH :query
        """
        
        params = {"user_id": user_id, "query": query}
        
        if doc_ids:
            placeholders = ','.join([f"'{doc_id}'" for doc_id in doc_ids])
            fts_query += f" AND doc_id IN ({placeholders})"
        
        fts_query += " LIMIT :limit"
        params["limit"] = limit
        
        result = db.execute(text(fts_query), params)
        rows = result.fetchall()
        
        return [
            {
                "doc_id": row[0],
                "page_number": row[1],
                "content": row[2],
                "snippet": row[3] if len(row) > 3 else row[2][:200]
            }
            for row in rows
        ]


def get_document_context(user_id: str, doc_ids: Optional[List[str]] = None, max_tokens: int = 2000) -> str:
    """
    Get document context for a user, optionally filtered by doc_ids.
    Returns concatenated content up to max_tokens.
    """
    create_tables()
    
    with SessionLocal() as db:
        query = db.query(Page).filter(Page.user_id == user_id)
        
        if doc_ids:
            query = query.filter(Page.doc_id.in_(doc_ids))
        
        pages = query.order_by(Page.doc_id, Page.page_number).all()
        
        context_parts = []
        total_tokens = 0
        
        for page in pages:
            if total_tokens + page.token_count > max_tokens:
                break
            
            context_parts.append(f"[Document: {page.doc_id}, Page {page.page_number}]\n{page.content}")
            total_tokens += page.token_count
        
        return "\n\n".join(context_parts)


